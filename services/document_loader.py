import os
import tempfile
from datetime import datetime
from typing import List
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from core.utils import normalize_text

def load_pdf(file_path: str, filename: str) -> List[Document]:
    loader = PDFPlumberLoader(file_path)
    docs = loader.load()
    upload_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    for doc in docs:
        page_num = doc.metadata.get("page", 0) + 1
        doc.metadata.update({
            "filename": filename, 
            "source": file_path,
            "page": page_num,
            "upload_date": upload_date,
            "file_type": "PDF"
        })
    return docs

def load_docx(file_path: str, filename: str) -> List[Document]:
    docx_file = DocxDocument(file_path)
    docs: List[Document] = []
    upload_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    for idx, para in enumerate(docx_file.paragraphs, start=1):
        text = normalize_text(para.text)
        if not text: continue
        docs.append(Document(
            page_content=text, 
            metadata={
                "filename": filename, 
                "paragraph_index": idx,
                "upload_date": upload_date,
                "file_type": "DOCX"
            }
        ))
    return docs

def split_documents_with_metadata(docs: List[Document], chunk_size: int, chunk_overlap: int) -> List[Document]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, 
        chunk_overlap=chunk_overlap
    )
    return text_splitter.split_documents(docs)

from concurrent.futures import ThreadPoolExecutor

def process_uploaded_files(uploaded_files, chunk_size=1000, chunk_overlap=100):
    all_raw_docs = []
    file_summaries = []

    with ThreadPoolExecutor() as executor:
        futures = []
        for uploaded_file in uploaded_files:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_file.name)[1]) as tmp_file:
                tmp_file.write(uploaded_file.getvalue())
                temp_path = tmp_file.name

            if uploaded_file.name.lower().endswith(".pdf"):
                futures.append(executor.submit(load_pdf, temp_path, uploaded_file.name))
            elif uploaded_file.name.lower().endswith(".docx"):
                futures.append(executor.submit(load_docx, temp_path, uploaded_file.name))

        for future in futures:
            try:
                docs = future.result()
                all_raw_docs.extend(docs)
                if docs:
                    file_summaries.append({
                        "filename": docs[0].metadata["filename"],
                        "type": "PDF" if ".pdf" in docs[0].metadata["filename"].lower() else "DOCX",
                        "raw_units": len(docs)
                    })
            except Exception as e:
                print(f"Error loading file: {e}")

    split_docs = split_documents_with_metadata(all_raw_docs, chunk_size, chunk_overlap)
    return all_raw_docs, split_docs, file_summaries
