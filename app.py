import os
import re
import tempfile
from typing import List, Tuple, Dict, Any

import streamlit as st
from docx import Document as DocxDocument

from langchain_core.documents import Document
from langchain_community.document_loaders import PDFPlumberLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama


# ==========================================
# CẤU HÌNH TRANG
# ==========================================
st.set_page_config(page_title="SmartDoc AI+", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #F8F9FA; }
    .stButton>button { border-radius: 8px; }
    .stFileUploader { border: 1px dashed #FFC107; padding: 10px; border-radius: 10px; }
    [data-testid="stSidebar"] { background-color: #2C2F33; color: #FFFFFF; }
    h1 { color: #212529; }
    .source-box {
        background: #f6f8fa;
        border-left: 4px solid #007BFF;
        padding: 12px;
        border-radius: 8px;
        margin-bottom: 10px;
        color: #111111;
    }
    mark {
        background-color: #fff59d;
        padding: 0.1em 0.2em;
        border-radius: 4px;
    }
    </style>
""", unsafe_allow_html=True)


# ==========================================
# SESSION STATE
# ==========================================
def init_session_state() -> None:
    defaults = {
        "vector_store": None,
        "retriever": None,
        "processed_docs": None,
        "raw_docs": None,
        "chat_history": [],
        "uploaded_filename": None,
        "uploaded_filetype": None,
        "confirm_clear_history": False,
        "confirm_clear_vector": False,
        "last_chunk_config": None,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_session_state()


# ==========================================
# HÀM PHỤ TRỢ
# ==========================================
def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def detect_vietnamese(question: str) -> bool:
    vietnamese_pattern = r"[ăâđêôơưáàảãạấầẩẫậắằẳẵặéèẻẽẹếềểễệíìỉĩịóòỏõọốồổỗộớờởỡợúùủũụứừửữựýỳỷỹỵ]"
    return bool(re.search(vietnamese_pattern, question.lower()))


def highlight_snippet(text: str, query: str, max_terms: int = 6) -> str:
    """
    Highlight các từ khóa trong câu hỏi xuất hiện trong source text.
    """
    safe_text = text
    words = re.findall(r"\w+", query.lower())
    words = [w for w in words if len(w) >= 3][:max_terms]

    for w in words:
        pattern = re.compile(rf"({re.escape(w)})", re.IGNORECASE)
        safe_text = pattern.sub(r"<mark>\1</mark>", safe_text)

    return safe_text


def clear_history() -> None:
    st.session_state.chat_history = []
    st.session_state.confirm_clear_history = False


def clear_vector_store() -> None:
    st.session_state.vector_store = None
    st.session_state.retriever = None
    st.session_state.processed_docs = None
    st.session_state.raw_docs = None
    st.session_state.uploaded_filename = None
    st.session_state.uploaded_filetype = None
    st.session_state.last_chunk_config = None
    st.session_state.confirm_clear_vector = False


def build_source_label(doc: Document) -> str:
    source_type = doc.metadata.get("source_type", "unknown").upper()

    if source_type == "PDF":
        page = doc.metadata.get("page", "N/A")
        chunk_id = doc.metadata.get("chunk_id", "N/A")
        start = doc.metadata.get("start_index", "N/A")
        end = doc.metadata.get("end_index", "N/A")
        return f"PDF | Trang {page} | Chunk {chunk_id} | Vị trí {start}-{end}"

    if source_type == "DOCX":
        para = doc.metadata.get("paragraph_index", "N/A")
        chunk_id = doc.metadata.get("chunk_id", "N/A")
        start = doc.metadata.get("start_index", "N/A")
        end = doc.metadata.get("end_index", "N/A")
        return f"DOCX | Đoạn {para} | Chunk {chunk_id} | Vị trí {start}-{end}"

    return "Nguồn không xác định"


def load_pdf(file_path: str) -> List[Document]:
    loader = PDFPlumberLoader(file_path)
    docs = loader.load()

    cleaned_docs = []
    for d in docs:
        page_num = d.metadata.get("page", None)
        if page_num is not None:
            page_num = page_num + 1  # PDFPlumber thường bắt đầu từ 0
        cleaned_docs.append(
            Document(
                page_content=normalize_text(d.page_content),
                metadata={
                    "source": file_path,
                    "source_type": "PDF",
                    "page": page_num if page_num is not None else "N/A",
                },
            )
        )
    return cleaned_docs


def load_docx(file_path: str) -> List[Document]:
    """
    Hỗ trợ DOCX bằng python-docx.
    Tách theo paragraph để giữ vị trí nguồn tốt hơn.
    """
    docx_file = DocxDocument(file_path)
    docs: List[Document] = []

    for idx, para in enumerate(docx_file.paragraphs, start=1):
        text = normalize_text(para.text)
        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "source_type": "DOCX",
                        "paragraph_index": idx,
                    },
                )
            )

    return docs


def split_documents_with_metadata(
    docs: List[Document],
    chunk_size: int,
    chunk_overlap: int,
) -> List[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        add_start_index=True,
    )

    split_docs = splitter.split_documents(docs)

    for i, d in enumerate(split_docs, start=1):
        d.metadata["chunk_id"] = i
        start_idx = d.metadata.get("start_index", 0)
        d.metadata["start_index"] = start_idx
        d.metadata["end_index"] = start_idx + len(d.page_content)

    return split_docs


@st.cache_resource(show_spinner=False)
def get_embedder():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-mpnet-base-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


@st.cache_resource(show_spinner=False)
def get_llm():
    return Ollama(
        model="qwen2.5:7b",
        temperature=0.7,
        top_p=0.9,
        repeat_penalty=1.1,
    )


def process_uploaded_file(
    uploaded_file,
    chunk_size: int,
    chunk_overlap: int,
) -> Tuple[List[Document], List[Document], FAISS]:
    suffix = os.path.splitext(uploaded_file.name)[1].lower()

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getbuffer())
        temp_path = tmp.name

    try:
        if suffix == ".pdf":
            raw_docs = load_pdf(temp_path)
        elif suffix == ".docx":
            raw_docs = load_docx(temp_path)
        else:
            raise ValueError("Chỉ hỗ trợ PDF và DOCX.")

        if not raw_docs:
            raise ValueError("Không trích xuất được nội dung từ tài liệu.")

        split_docs = split_documents_with_metadata(
            raw_docs,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

        embedder = get_embedder()
        vector_store = FAISS.from_documents(split_docs, embedder)

        return raw_docs, split_docs, vector_store
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)


def build_prompt(context: str, question: str) -> str:
    if detect_vietnamese(question):
        return f"""Bạn là trợ lý hỏi đáp tài liệu.
Chỉ trả lời dựa trên ngữ cảnh được cung cấp.
Nếu thông tin không đủ, hãy nói rõ là không tìm thấy trong tài liệu.
Trả lời ngắn gọn, chính xác, bằng tiếng Việt.

Ngữ cảnh:
{context}

Câu hỏi:
{question}

Trả lời:"""

    return f"""You are a document QA assistant.
Answer only based on the given context.
If the answer is not in the document, say so clearly.
Keep the answer concise and accurate.

Context:
{context}

Question:
{question}

Answer:"""


def compare_chunk_strategy_guidance(chunk_size: int, chunk_overlap: int) -> str:
    """
    Gợi ý báo cáo định tính để phục vụ mục 8.2.4.
    Không bịa số liệu đo đạc.
    """
    if chunk_size == 500:
        size_note = "Chunk nhỏ: truy xuất chính xác chi tiết tốt hơn, nhưng dễ thiếu ngữ cảnh tổng thể."
    elif chunk_size == 1000:
        size_note = "Chunk cân bằng: thường là mức ổn cho cả độ chính xác và tốc độ."
    elif chunk_size == 1500:
        size_note = "Chunk lớn hơn: giữ được nhiều ngữ cảnh hơn, nhưng đôi khi truy xuất kém tập trung."
    else:
        size_note = "Chunk rất lớn: ngữ cảnh rộng, nhưng có thể kéo theo thông tin thừa."

    if chunk_overlap == 50:
        overlap_note = "Overlap thấp: nhanh hơn, nhưng dễ mất ý ở ranh giới chunk."
    elif chunk_overlap == 100:
        overlap_note = "Overlap vừa phải: thường là lựa chọn an toàn."
    else:
        overlap_note = "Overlap cao: giữ mạch nội dung tốt hơn, nhưng tăng trùng lặp."

    return f"{size_note} {overlap_note}"


# ==========================================
# SIDEBAR
# ==========================================
st.sidebar.title("⚙️ Cấu hình hệ thống")
st.sidebar.markdown("---")
st.sidebar.write("**Model:** Qwen2.5:7b")
st.sidebar.write("**Embeddings:** Multilingual MPNet")
st.sidebar.write("**Status:** Local Runtime (Ollama)")

st.sidebar.subheader("Chunk Strategy")
chunk_size = st.sidebar.selectbox(
    "Chunk size",
    options=[500, 1000, 1500, 2000],
    index=1,
)
chunk_overlap = st.sidebar.selectbox(
    "Chunk overlap",
    options=[50, 100, 200],
    index=1,
)

st.sidebar.caption(compare_chunk_strategy_guidance(chunk_size, chunk_overlap))

st.sidebar.markdown("---")
st.sidebar.subheader("Lịch sử hội thoại")

if st.session_state.chat_history:
    for idx, item in enumerate(st.session_state.chat_history, start=1):
        with st.sidebar.expander(f"Câu {idx}: {item['question'][:35]}...", expanded=False):
            st.markdown(f"**Hỏi:** {item['question']}")
            st.markdown(f"**Đáp:** {item['answer']}")
            st.caption(
                f"File: {item.get('filename', 'N/A')} | "
                f"Chunk: {item.get('chunk_size', 'N/A')}/{item.get('chunk_overlap', 'N/A')}"
            )
else:
    st.sidebar.info("Chưa có lịch sử chat.")

st.sidebar.markdown("---")

col_h1, col_h2 = st.sidebar.columns(2)

with col_h1:
    if st.button("Clear History", use_container_width=True):
        st.session_state.confirm_clear_history = True

with col_h2:
    if st.button("Clear Vector Store", use_container_width=True):
        st.session_state.confirm_clear_vector = True

if st.session_state.confirm_clear_history:
    st.sidebar.warning("Bạn có chắc muốn xóa toàn bộ lịch sử chat?")
    c1, c2 = st.sidebar.columns(2)
    with c1:
        if st.button("Xác nhận xóa lịch sử", key="confirm_history_delete", use_container_width=True):
            clear_history()
            st.sidebar.success("Đã xóa lịch sử chat.")
    with c2:
        if st.button("Hủy", key="cancel_history_delete", use_container_width=True):
            st.session_state.confirm_clear_history = False

if st.session_state.confirm_clear_vector:
    st.sidebar.warning("Bạn có chắc muốn xóa vector store và tài liệu hiện tại?")
    c1, c2 = st.sidebar.columns(2)
    with c1:
        if st.button("Xác nhận xóa tài liệu", key="confirm_vector_delete", use_container_width=True):
            clear_vector_store()
            st.sidebar.success("Đã xóa vector store.")
    with c2:
        if st.button("Hủy ", key="cancel_vector_delete", use_container_width=True):
            st.session_state.confirm_clear_vector = False


# ==========================================
# GIAO DIỆN CHÍNH
# ==========================================
st.title("📄 SmartDoc AI+ - Intelligent Document Q&A System")
st.write("Hệ thống hỏi đáp tài liệu thông minh hỗ trợ PDF và DOCX, có lưu lịch sử hội thoại, citation và tùy chỉnh chunk strategy.")

uploaded_file = st.file_uploader(
    "Tải lên tài liệu để bắt đầu",
    type=["pdf", "docx"],
    help="Hỗ trợ cả PDF và DOCX",
)

# ==========================================
# XỬ LÝ FILE
# ==========================================
if uploaded_file is not None:
    need_reprocess = (
        st.session_state.vector_store is None
        or st.session_state.uploaded_filename != uploaded_file.name
        or st.session_state.last_chunk_config != (chunk_size, chunk_overlap)
    )

    if need_reprocess:
        with st.status("Đang xử lý tài liệu...", expanded=True) as status:
            st.write("1. Đang đọc nội dung tài liệu...")
            st.write("2. Đang tách chunk theo tham số người dùng...")
            st.write("3. Đang tạo embeddings...")
            st.write("4. Đang khởi tạo FAISS vector store...")

            try:
                raw_docs, split_docs, vector_store = process_uploaded_file(
                    uploaded_file,
                    chunk_size=chunk_size,
                    chunk_overlap=chunk_overlap,
                )

                st.session_state.raw_docs = raw_docs
                st.session_state.processed_docs = split_docs
                st.session_state.vector_store = vector_store
                st.session_state.retriever = vector_store.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 3, "fetch_k": 20},
                )
                st.session_state.uploaded_filename = uploaded_file.name
                st.session_state.uploaded_filetype = os.path.splitext(uploaded_file.name)[1].lower()
                st.session_state.last_chunk_config = (chunk_size, chunk_overlap)

                status.update(label="Tài liệu đã sẵn sàng.", state="complete", expanded=False)

            except Exception as e:
                status.update(label="Xử lý tài liệu thất bại.", state="error", expanded=True)
                st.error(f"Lỗi khi xử lý tài liệu: {e}")

    st.success(
        f"Đã nạp tài liệu: {st.session_state.uploaded_filename} | "
        f"Chunk size = {chunk_size}, overlap = {chunk_overlap}"
    )

    user_question = st.text_input("💬 Đặt câu hỏi về nội dung tài liệu của bạn:")

    if user_question and st.session_state.retriever is not None:
        with st.spinner("Đang suy nghĩ..."):
            try:
                relevant_docs = st.session_state.retriever.invoke(user_question)
                context = "\n\n".join([doc.page_content for doc in relevant_docs])

                llm = get_llm()
                prompt = build_prompt(context, user_question)
                response = llm.invoke(prompt)

                # Lưu lịch sử hội thoại trong session
                st.session_state.chat_history.append(
                    {
                        "question": user_question,
                        "answer": response,
                        "filename": st.session_state.uploaded_filename,
                        "chunk_size": chunk_size,
                        "chunk_overlap": chunk_overlap,
                        "sources": [
                            {
                                "label": build_source_label(doc),
                                "content": doc.page_content,
                                "metadata": doc.metadata,
                            }
                            for doc in relevant_docs
                        ],
                    }
                )

                st.subheader("💡 Trả lời")
                st.write(response)

                st.subheader("📚 Citation / Source Tracking")
                for i, doc in enumerate(relevant_docs, start=1):
                    source_label = build_source_label(doc)
                    highlighted = highlight_snippet(doc.page_content, user_question)

                    with st.expander(f"Nguồn {i}: {source_label}", expanded=False):
                        st.markdown(
                            f"""
                            <div class="source-box">
                                {highlighted}
                            </div>
                            """,
                            unsafe_allow_html=True,
                        )
                        st.caption("Đã highlight các từ khóa liên quan từ câu hỏi trong đoạn ngữ cảnh gốc.")

            except Exception as e:
                st.error(f"Lỗi khi sinh câu trả lời: {e}")

    # Hiển thị hội thoại hiện tại ở main area
    if st.session_state.chat_history:
        st.markdown("---")
        st.subheader("🕘 Lịch sử hỏi đáp trong phiên hiện tại")

        for idx, item in enumerate(reversed(st.session_state.chat_history), start=1):
            with st.expander(f"Lượt gần đây {idx}: {item['question']}", expanded=False):
                st.markdown(f"**Hỏi:** {item['question']}")
                st.markdown(f"**Đáp:** {item['answer']}")
                st.caption(
                    f"File: {item.get('filename', 'N/A')} | "
                    f"Chunk size: {item.get('chunk_size', 'N/A')} | "
                    f"Overlap: {item.get('chunk_overlap', 'N/A')}"
                )

                if item.get("sources"):
                    st.markdown("**Nguồn đã dùng:**")
                    for s_idx, src in enumerate(item["sources"], start=1):
                        with st.expander(f"Xem context nguồn {s_idx}: {src['label']}", expanded=False):
                            st.write(src["content"])

    st.markdown("---")
    st.subheader("📊 Gợi ý báo cáo cho mục 8.2.4")
    st.write(
        f"""
        Cấu hình hiện tại:
        - Chunk size: **{chunk_size}**
        - Chunk overlap: **{chunk_overlap}**

        Nhận xét định tính:
        - {compare_chunk_strategy_guidance(chunk_size, chunk_overlap)}

        Để so sánh độ chính xác, bạn nên:
        1. Chuẩn bị một bộ 10-20 câu hỏi cố định.
        2. Chạy lần lượt với các cấu hình:
           - chunk_size: 500, 1000, 1500, 2000
           - chunk_overlap: 50, 100, 200
        3. Ghi nhận:
           - câu trả lời đúng/sai
           - có đủ ý không
           - tốc độ phản hồi
           - mức độ liên quan của nguồn trích dẫn
        4. Chọn cấu hình cân bằng nhất giữa độ chính xác và tốc độ.
        """
    )

else:
    st.info("Vui lòng tải lên một file PDF hoặc DOCX để bắt đầu trò chuyện với tài liệu.")